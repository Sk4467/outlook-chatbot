import imaplib
import email
from email.header import decode_header
import io
from typing import List, Dict, Any, Optional

from bs4 import BeautifulSoup
import chardet


def _decode_mime_words(s: Optional[str]) -> str:
    if not s:
        return ""
    try:
        dh = decode_header(s)
        parts = []
        for p, enc in dh:
            if isinstance(p, bytes):
                try:
                    parts.append(p.decode(enc or "utf-8", errors="ignore"))
                except Exception:
                    parts.append(p.decode("utf-8", errors="ignore"))
            else:
                parts.append(p)
        return "".join(parts)
    except Exception:
        return s


def _html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator=" ", strip=True)


def _part_to_text(part: email.message.Message) -> str:
    ctype = (part.get_content_type() or "").lower()
    charset = part.get_content_charset() or "utf-8"
    payload = part.get_payload(decode=True) or b""
    if not payload:
        return ""
    try:
        text = payload.decode(charset, errors="ignore")
    except Exception:
        try:
            enc = chardet.detect(payload).get("encoding") or "utf-8"
            text = payload.decode(enc, errors="ignore")
        except Exception:
            text = payload.decode("utf-8", errors="ignore")
    if ctype == "text/html":
        return _html_to_text(text)
    return text


def _parse_pdf_bytes(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        out = []
        for i in range(len(doc)):
            out.append(doc[i].get_text("text"))
        doc.close()
        return "\n".join(t for t in out if t)
    except Exception as e:
        return f"[PDF parse error: {e}]"


def _parse_docx_bytes(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        return f"[DOCX parse error: {e}]"


def _parse_excel_like_bytes(data: bytes, filename: str) -> str:
    # Handles .xlsx, .xlsm, .csv
    try:
        fn = (filename or "").lower()
        if fn.endswith(".csv"):
            raw = data
            enc = "utf-8-sig"
            try:
                det = chardet.detect(raw)
                if det and det.get("encoding"):
                    enc = det["encoding"]
            except Exception:
                pass
            try:
                txt = raw.decode(enc, errors="ignore")
            except Exception:
                txt = raw.decode("utf-8", errors="ignore")
            lines = txt.splitlines()
            return "\n".join(lines[:200])
        else:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(data), data_only=True)
            out = []
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                headers = [(str(h).strip() if h is not None else "") for h in rows[0]]
                for r_idx, row in enumerate(rows[1:], start=2):
                    parts = []
                    for c_idx, val in enumerate(row):
                        header = headers[c_idx] if c_idx < len(headers) else f"Col{c_idx+1}"
                        val_str = "" if val is None else str(val)
                        parts.append(f"{header}: {val_str}")
                    out.append(f"Sheet: {ws.title}, Row {r_idx} | " + " | ".join(parts))
            return "\n".join(out)
    except Exception as e:
        return f"[Excel parse error: {e}]"


def _parse_attachment_bytes(filename: str, content_type: str, data: bytes) -> str:
    fn = (filename or "").lower()
    ct = (content_type or "").lower()
    if fn.endswith(".pdf") or "pdf" in ct:
        return _parse_pdf_bytes(data)
    if fn.endswith(".docx") or "officedocument.wordprocessingml.document" in ct:
        return _parse_docx_bytes(data)
    if fn.endswith((".xlsx", ".xlsm", ".csv")) or "spreadsheetml" in ct or ct == "text/csv":
        return _parse_excel_like_bytes(data, filename or "file")
    return f"[Unsupported attachment type: {filename or content_type}]"


def fetch_emails(
    email_account: str,
    email_password: str,
    n: int = 10,
    folder: str = "INBOX",
    include_attachments: bool = True,
    max_attachment_bytes: int = 15 * 1024 * 1024,
) -> List[Dict[str, Any]]:
    imap = imaplib.IMAP4_SSL("imap.gmail.com")
    try:
        imap.login(email_account, email_password)
        status, _ = imap.select(folder)
        if status != "OK":
            raise RuntimeError(f"Cannot select folder {folder}")

        status, data = imap.search(None, "ALL")
        if status != "OK" or not data or not data[0]:
            return []

        all_ids = data[0].split()
        target_ids = all_ids if (n is None or n <= 0) else all_ids[-n:]

        results: List[Dict[str, Any]] = []

        for eid in target_ids:
            status, msg_data = imap.fetch(eid, "(RFC822)")
            if status != "OK" or not msg_data or not msg_data[0]:
                continue
            raw = msg_data[0][1]
            msg = email.message_from_bytes(raw)

            subject = _decode_mime_words(msg.get("Subject"))
            from_ = _decode_mime_words(msg.get("From"))
            date_ = _decode_mime_words(msg.get("Date"))

            body_text = ""
            attachments_info = []

            if msg.is_multipart():
                plain_found = False
                html_fallback = None
                for part in msg.walk():
                    disp = str(part.get("Content-Disposition") or "").lower()
                    ctype = (part.get_content_type() or "").lower()
                    if ctype == "text/plain" and "attachment" not in disp:
                        text = _part_to_text(part)
                        if text:
                            body_text += (("\n" if body_text else "") + text)
                            plain_found = True
                    elif ctype == "text/html" and "attachment" not in disp:
                        if html_fallback is None:
                            html_fallback = _part_to_text(part)
                if not plain_found and html_fallback:
                    body_text = html_fallback

                if include_attachments:
                    for part in msg.walk():
                        disp = str(part.get("Content-Disposition") or "").lower()
                        if ("attachment" in disp) or (part.get_filename() is not None):
                            filename = _decode_mime_words(part.get_filename() or "")
                            ctype = (part.get_content_type() or "").lower()
                            data = part.get_payload(decode=True) or b""
                            size = len(data)
                            if size == 0:
                                continue
                            if size > max_attachment_bytes:
                                attachments_info.append({
                                    "filename": filename,
                                    "content_type": ctype,
                                    "size": size,
                                    "text": f"[Skipped: {size} bytes > max {max_attachment_bytes}]",
                                })
                                continue
                            text = _parse_attachment_bytes(filename, ctype, data)
                            attachments_info.append({
                                "filename": filename,
                                "content_type": ctype,
                                "size": size,
                                "text": text,
                            })
            else:
                ctype = (msg.get_content_type() or "").lower()
                if ctype in ("text/plain", "text/html"):
                    body_text = _part_to_text(msg)
                else:
                    payload = msg.get_payload(decode=True) or b""
                    body_text = f"[Non-text body: {ctype}, {len(payload)} bytes]"

            attach_concat = ""
            for a in attachments_info:
                label = a["filename"] or a["content_type"]
                attach_concat += f"\n\n[Attachment: {label}]\n{a['text']}"

            combined = (
                f"From: {from_}\nSubject: {subject}\nDate: {date_}\n\n{body_text}{attach_concat}"
            )

            results.append({
                "id": eid.decode("ascii", errors="ignore"),
                "from": from_,
                "subject": subject,
                "date": date_,
                "body_text": body_text,
                "attachments": attachments_info,
                "combined_text": combined,
            })

        return results
    finally:
        try:
            imap.close()
        except Exception:
            pass
        try:
            imap.logout()
        except Exception:
            pass

